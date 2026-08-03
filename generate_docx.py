"""Generate a professional Word document - black text, Calibri font, no background colors."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

BLACK = RGBColor(0x00, 0x00, 0x00)
FONT = "Calibri"


def sf(run, size=Pt(10), bold=False, italic=False):
    """Set font to Calibri, black."""
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = FONT
        r.font.color.rgb = BLACK
    return h


def para(doc, text, size=Pt(10), bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    sf(r, size, bold, italic)
    return p


def bullet(doc, text, indent=Cm(1.5)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("\u2022  ")
    sf(r, Pt(9), bold=False)
    r2 = p.add_run(text)
    sf(r2, Pt(9))


def shade(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


doc = Document()

# Default style
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = Pt(10)
style.font.color.rgb = BLACK

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# TITLE
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("UHHomes")
sf(r, Pt(28), bold=True)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Role-Based Architecture \u2014 All User Accounts, Flows & Feature Assignments")
sf(r, Pt(11), italic=True)

doc.add_paragraph()

# ===================== SECTION 1 =====================
heading(doc, "1. All Possible User Roles in the Platform")
para(doc, "A mature home construction platform needs these 5 distinct roles, each with clear responsibilities.",
     Pt(10), italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

roles = [
    ("Homebuyer / Owner", "End Customer",
     "The person who purchased the home. Tracks construction, views updates, asks questions, downloads documents.",
     ["View project dashboard & progress", "Browse construction timeline",
      "View photo/video gallery", "Submit inquiries to builder",
      "Receive alerts & notifications", "Download documents (contracts, permits)",
      "Favorite properties for comparison", "Receive AI weekly digest email"]),
    ("Company Admin", "Super Administrator",
     "The construction company owner or top-level manager. Has full access to everything \u2014 creates roles, manages the platform.",
     ["Full CRUD on all entities", "Create & manage user accounts",
      "Assign roles (PM, Sales, Architect)", "View company-wide analytics",
      "Generate PDF reports", "Manage properties catalog",
      "Broadcast alerts to all buyers", "Configure platform settings",
      "Trigger AI summary generation"]),
    ("Project Manager", "Projects & Site",
     "Assigned to one or more projects. Oversees milestones, coordinates with buyers, ensures timelines are met, and handles on-ground site supervision.",
     ["View & update assigned projects", "Create & manage milestones",
      "Post daily/weekly site updates", "Upload photos & videos from site",
      "Update milestone progress %", "Log issues & delays",
      "Respond to buyer inquiries", "Generate project-level reports",
      "Manage project documents", "Escalate delays to Admin",
      "Inspect & approve milestones"]),
    ("Sales Agent / Realtor", "Customer Acquisition",
     "Handles property inquiries from prospects, onboards new buyers, manages the pre-sale relationship until project assignment.",
     ["View property catalog", "Manage prospect leads",
      "Schedule property tours", "Onboard new buyer accounts",
      "Track inquiry-to-sale pipeline", "View assigned buyer profiles",
      "No access to construction data"]),
    ("Architect / Designer", "Design & Plans",
     "Uploads floor plans, elevations, 3D renders. Handles design-related buyer requests like customizations or material selections.",
     ["Upload & manage floor plans", "Upload elevation renders",
      "Handle design change requests", "View assigned projects",
      "Comment on milestone updates", "No access to financial data"]),
]

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = tbl.rows[0].cells
for i, txt in enumerate(["ROLE", "DESCRIPTION", "KEY RESPONSIBILITIES"]):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    r = p.add_run(txt)
    sf(r, Pt(9), bold=True)
    shade(hdr[i], "F2F2F2")

for name, tag, desc, resps in roles:
    row = tbl.add_row().cells
    row[0].text = ""
    p = row[0].paragraphs[0]
    r = p.add_run(name)
    sf(r, Pt(10), bold=True)
    p2 = row[0].add_paragraph()
    r2 = p2.add_run(tag)
    sf(r2, Pt(8), italic=True)

    row[1].text = ""
    p = row[1].paragraphs[0]
    r = p.add_run(desc)
    sf(r, Pt(9))

    row[2].text = ""
    for j, resp in enumerate(resps):
        p = row[2].paragraphs[0] if j == 0 else row[2].add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"\u2022 {resp}")
        sf(r, Pt(8))

for row in tbl.rows:
    row.cells[0].width = Cm(4.5)
    row.cells[1].width = Cm(6)
    row.cells[2].width = Cm(7)

doc.add_paragraph()

# ===================== SECTION 2 =====================
heading(doc, "2. Role Hierarchy & Access Control")
para(doc, "Who reports to whom, and who can see what.", Pt(10), italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

hierarchy = [
    ("Company Admin", "Full Access \u2014 Top Level"),
    ("    \u251c\u2500\u2500 Project Manager", "Project-Scoped + Site Supervision"),
    ("    \u2502       \u2514\u2500\u2500 Architect", "Designs Only (reports to PM)"),
    ("    \u251c\u2500\u2500 Sales Agent", "Pre-Sale Only"),
    ("    \u2514\u2500\u2500 Homebuyer", "Read + Inquire (consumer)"),
]
for item, desc in hierarchy:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(item)
    sf(r, Pt(11), bold=True)
    r2 = p.add_run(f"  \u2014  {desc}")
    sf(r2, Pt(9), italic=True)

doc.add_paragraph()
para(doc, "Relationships:", Pt(10), bold=True)
for rel in [
    "Admin \u2192 assigns projects \u2192 Project Manager",
    "Admin \u2192 assigns leads \u2192 Sales Agent",
    "Project Manager \u2192 requests designs \u2192 Architect",
    "Project Manager \u2192 posts updates visible to \u2192 Homebuyer",
    "Project Manager \u2192 responds to inquiries from \u2192 Homebuyer",
    "Sales Agent \u2192 onboards \u2192 Homebuyer",
]:
    bullet(doc, rel)

doc.add_paragraph()

# ===================== SECTION 3 =====================
heading(doc, "3. Complete Lifecycle \u2014 From Land to Handover")
para(doc, "Every step of a home construction project, mapped to which role acts and which role consumes.",
     Pt(10), italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

stages = [
    ("STAGE 1 \u2014 LEAD & SALE", [
        "Sales Agent shows property catalog & schedules tour for Buyer",
        "Buyer expresses interest / inquires",
        "Sales Agent converts lead to buyer account via Admin",
        "Admin sends welcome email + portal access",
        "Admin sends booking payment invoice",
        "Buyer makes booking payment",
    ]),
    ("STAGE 2 \u2014 PROJECT INITIATION", [
        "Admin creates project, assigns PM",
        "PM requests floor plans & design from Architect",
        "Architect uploads floor plans & elevations",
        "PM defines milestones & timeline",
        "System notifies Buyer: project has started",
    ]),
    ("STAGE 3 \u2014 CONSTRUCTION (Repeats per Milestone)", [
        "PM posts daily update + photos from site",
        "System notifies Buyer of progress update",
        "Buyer views timeline, gallery, updates",
        "Buyer submits question / concern",
        "PM responds to inquiry",
        "PM marks milestone 100%",
        "PM inspects & approves milestone",
        "Admin sends milestone payment invoice",
        "Buyer makes payment",
        "System auto-calculates overall progress %",
    ]),
    ("STAGE 4 \u2014 AUTOMATED SERVICES (Ongoing)", [
        "System generates AI weekly summary (CRON)",
        "System sends weekly progress digest email to Buyer",
        "System sends delayed milestone alert to PM",
        "System sends monthly analytics report to Admin",
    ]),
    ("STAGE 5 \u2014 HANDOVER & CLOSE", [
        "PM marks final milestone complete",
        "PM final inspection approved",
        "PM uploads completion certificate & warranty documents",
        "Admin sends final payment invoice",
        "Buyer makes final payment",
        "Admin marks project COMPLETED",
        "System notifies Buyer: home is ready",
        "Buyer downloads report, documents, gallery archive",
    ]),
]

for stage_name, steps in stages:
    para(doc, stage_name, Pt(11), bold=True)
    for step in steps:
        bullet(doc, step)
    doc.add_paragraph()

# ===================== SECTION 4 =====================
heading(doc, "4. Feature Access Matrix \u2014 Who Gets What")
para(doc, "\u25cf = Full Access    \u25d0 = Limited    \u25cb = No Access",
     Pt(9), italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

features = [
    ["View Dashboard",            "\u25cf", "\u25cf", "\u25cf", "\u25d0", "\u25cb"],
    ["Manage Users",              "\u25cb", "\u25cf", "\u25cb", "\u25cb", "\u25cb"],
    ["Create Projects",           "\u25cb", "\u25cf", "\u25cb", "\u25cb", "\u25cb"],
    ["Update Milestones",         "\u25cb", "\u25cf", "\u25cf", "\u25cb", "\u25cb"],
    ["Post Site Updates",         "\u25cb", "\u25cf", "\u25cf", "\u25cb", "\u25cb"],
    ["Upload Photos/Media",       "\u25cb", "\u25cf", "\u25cf", "\u25cb", "\u25cf"],
    ["View Construction Progress","\u25cf", "\u25cf", "\u25cf", "\u25cb", "\u25d0"],
    ["Submit Inquiries",          "\u25cf", "\u25cb", "\u25cb", "\u25cb", "\u25cb"],
    ["Respond to Inquiries",      "\u25cb", "\u25cf", "\u25cf", "\u25cb", "\u25cb"],
    ["Upload Floor Plans",        "\u25cb", "\u25cf", "\u25cb", "\u25cb", "\u25cf"],
    ["Manage Properties",         "\u25cb", "\u25cf", "\u25cb", "\u25d0", "\u25cb"],
    ["Manage Leads / Prospects",  "\u25cb", "\u25cf", "\u25cb", "\u25cf", "\u25cb"],
    ["Inspect & Approve Milestones","\u25cb", "\u25cf", "\u25cf", "\u25cb", "\u25cb"],
    ["Manage Invoices & Payments","\u25d0", "\u25cf", "\u25cb", "\u25cb", "\u25cb"],
    ["View/Download Documents",   "\u25cf", "\u25cf", "\u25cf", "\u25cb", "\u25d0"],
    ["Generate Reports (PDF)",    "\u25cb", "\u25cf", "\u25cf", "\u25cb", "\u25cb"],
    ["Broadcast Alerts",          "\u25cb", "\u25cf", "\u25d0", "\u25cb", "\u25cb"],
    ["AI Summary Generation",     "\u25cb", "\u25cf", "\u25cf", "\u25cb", "\u25cb"],
    ["Favorite Properties",       "\u25cf", "\u25cb", "\u25cb", "\u25cb", "\u25cb"],
]

headers = ["Feature", "Buyer", "Admin", "PM", "Sales", "Architect"]
ftbl = doc.add_table(rows=1, cols=6)
ftbl.style = 'Table Grid'
ftbl.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = ftbl.rows[0].cells
for i, txt in enumerate(headers):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(txt)
    sf(r, Pt(9), bold=True)
    shade(hdr[i], "F2F2F2")

for row_data in features:
    row = ftbl.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = ""
        p = row[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(val)
        sf(r, Pt(9))

for row in ftbl.rows:
    row.cells[0].width = Cm(5)
    for i in range(1, 6):
        row.cells[i].width = Cm(2.5)

doc.add_paragraph()

# ===================== SECTION 5 =====================
heading(doc, "5. How Data Flows Between Roles")
para(doc, "Every arrow represents a real interaction \u2014 who creates the data, who consumes it.",
     Pt(10), italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

flows = [
    ("PRE-SALE PHASE", [
        "Sales Agent \u2192 shows properties, schedules tours \u2192 Prospect",
        "Prospect \u2192 submits inquiry, makes booking \u2192 Sales Agent",
        "Sales Agent \u2192 converts to buyer account \u2192 Admin",
    ]),
    ("CONSTRUCTION PHASE", [
        "Architect \u2192 floor plans + designs \u2192 Project Manager",
        "Project Manager \u2192 site updates, photos & milestones \u2192 System",
        "Project Manager \u2192 inspects & approves \u2192 System",
    ]),
    ("BUYER EXPERIENCE", [
        "Buyer \u2192 views dashboard, timeline, gallery \u2192 Portal",
        "Buyer \u2192 asks questions \u2192 Inquiries",
        "Portal \u2192 real-time updates \u2192 Buyer",
        "AI Digest \u2192 weekly email \u2192 Buyer",
    ]),
]

for phase_name, items in flows:
    para(doc, phase_name, Pt(11), bold=True)
    for item in items:
        bullet(doc, item)
    doc.add_paragraph()

# ===================== SECTION 6 =====================
heading(doc, "6. When to Introduce Each Role (Development Roadmap)")
para(doc, "Phased approach \u2014 start simple, expand as the platform grows.",
     Pt(10), italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

roadmap = [
    ("PHASE 1 \u2014 MVP", "Months 1\u20133 \u00b7 Core platform with 2 roles", [
        ("Company Admin", "Full platform access \u2014 create properties, projects, manage users", "Done"),
        ("Homebuyer", "Track construction, view timeline, gallery, submit inquiries", "Done"),
    ]),
    ("PHASE 2 \u2014 GROWTH", "Months 4\u20137 \u00b7 Delegate project & site work", [
        ("Project Manager", "Manage projects, milestones, site updates, photos, inspections, buyer inquiries", "In Progress"),
    ]),
    ("PHASE 3 \u2014 SCALE", "Months 7\u201310 \u00b7 Pre-sale pipeline & design workflows", [
        ("Sales Agent / Realtor", "Manage prospect leads, schedule tours, onboard new buyers", "Planned"),
        ("Architect / Designer", "Upload floor plans, elevations, 3D renders; handle design change requests", "Planned"),
    ]),
    ("PHASE 4 \u2014 ENTERPRISE", "Months 10\u201314 \u00b7 Multi-builder support", [
        ("Multi-Tenant", "Support multiple construction companies on one platform", "Future"),
    ]),
]

for phase_name, sub, items in roadmap:
    para(doc, phase_name, Pt(12), bold=True)
    para(doc, sub, Pt(9), italic=True)
    for name, desc, status in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{name}  ")
        sf(r, Pt(10), bold=True)
        r2 = p.add_run(f"[{status}]")
        sf(r2, Pt(8), italic=True)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1.5)
        r3 = p2.add_run(desc)
        sf(r3, Pt(9))
    doc.add_paragraph()

# ===================== SECTION 7 =====================
heading(doc, "7. Notification Flow \u2014 Who Notifies Whom")
para(doc, "Every notification mapped: what triggers it, who receives it.",
     Pt(10), italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

para(doc, "EVENT TRIGGERS", Pt(10), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

triggers = ["Site update posted", "Milestone completed", "Invoice sent", "Inquiry received",
            "Payment received", "Project assigned", "Weekly cron job", "Delay detected"]

ttbl = doc.add_table(rows=2, cols=4)
ttbl.style = 'Table Grid'
ttbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, t in enumerate(triggers):
    cell = ttbl.rows[i // 4].cells[i % 4]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t)
    sf(r, Pt(9))

doc.add_paragraph()
para(doc, "\u2193", Pt(14), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "NOTIFICATION ENGINE", Pt(11), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "System Router \u2014 decides who gets what", Pt(9), italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "\u2193", Pt(14), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

para(doc, "RECIPIENTS", Pt(10), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

recipients = [
    ("Homebuyer", "EMAIL \u00b7 PUSH \u00b7 IN-APP", [
        "Construction updates", "Invoice & payment receipts",
        "Inquiry replies", "Weekly AI digest", "Project milestones"
    ]),
    ("Project Manager", "EMAIL \u00b7 IN-APP \u00b7 PUSH", [
        "Delayed milestone alerts", "Buyer inquiries",
        "New project assignments", "Payment confirmations", "Escalation requests"
    ]),
    ("Company Admin", "EMAIL \u00b7 DASHBOARD", [
        "Weekly analytics report", "Escalations from PM",
        "Payment received alerts", "New user registrations", "System health warnings"
    ]),
]

rtbl = doc.add_table(rows=1, cols=3)
rtbl.style = 'Table Grid'
rtbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (name, channels, items) in enumerate(recipients):
    cell = rtbl.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(name)
    sf(r, Pt(11), bold=True)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(channels)
    sf(r2, Pt(8), italic=True)
    for item in items:
        p3 = cell.add_paragraph()
        p3.paragraph_format.space_after = Pt(2)
        r3 = p3.add_run(f"\u2022 {item}")
        sf(r3, Pt(8))

doc.add_paragraph()
doc.add_paragraph()

# Footer
ft = doc.add_paragraph()
ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ft.add_run("UHHomes Role-Based Architecture \u2022 Generated for development planning")
sf(r, Pt(8), italic=True)

# Save
output = "/Users/sowjanyamalipeddi/CascadeProjects/windsurf-project/UHHomes_Role_Architecture.docx"
doc.save(output)
print(f"Done: {output}")
